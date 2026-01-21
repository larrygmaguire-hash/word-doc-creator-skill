#!/usr/bin/env python3
"""
Client Document Generator
=========================
Creates professional Word documents from markdown source using your letterhead template.

CUSTOMISATION REQUIRED:
- Update YOUR_NAME, YOUR_QUALIFICATIONS, YOUR_TITLE in the sign-off section
- Adjust font settings (FONT_NAME, BODY_FONT_SIZE, etc.) to match your brand
- Modify spacing values to suit your document style

Requirements:
- python-docx library (pip install python-docx)
- A Word document letterhead template with headers/footers

Usage:
    python create_client_doc.py --template "path/to/letterhead.docx" \
        --source "path/to/content.md" \
        --output "path/to/output.docx" \
        --recipient-name "John Smith" \
        --recipient-title "Director" \
        --recipient-org "Acme Corp" \
        --recipient-address "123 Main Street" \
        --recipient-city "Dublin 1" \
        --recipient-country "Ireland" \
        --doc-title "Project Proposal" \
        --date "Monday 20th January 2026"
"""

import argparse
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# =============================================================================
# CUSTOMISE THESE SETTINGS FOR YOUR BRAND
# =============================================================================

# Font settings
FONT_NAME = 'Calibri'           # Change to your preferred font
BODY_FONT_SIZE = 11             # Body text size in points
QUALIFICATIONS_FONT_SIZE = 10   # Smaller size for qualifications line
HEADING_2_SIZE = 14             # H2 heading size
HEADING_3_SIZE = 12             # H3 heading size

# Line spacing (1.0 = single, 1.5 = one-and-a-half, 2.0 = double)
LINE_SPACING = 1.2

# Your sign-off details - CHANGE THESE
YOUR_NAME = "Your Name"
YOUR_QUALIFICATIONS = "Your Qualifications Here"
YOUR_TITLE = "Your Professional Title"

# =============================================================================
# END OF CUSTOMISATION SECTION
# =============================================================================


def parse_markdown(md_path: Path) -> dict:
    """Parse markdown file into structured sections."""
    content = md_path.read_text(encoding='utf-8')

    # Split on page break marker
    if '\\newpage' in content:
        cover_letter, remaining = content.split('\\newpage', 1)
    else:
        cover_letter = content
        remaining = ""

    return {
        'cover_letter': cover_letter.strip(),
        'content_sections': remaining.strip()
    }


def set_paragraph_spacing(paragraph, space_before=None, space_after=None, line_spacing=LINE_SPACING):
    """Set paragraph spacing."""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')

    if space_before is not None:
        spacing.set(qn('w:before'), str(int(space_before * 20)))
    else:
        spacing.set(qn('w:before'), '0')

    if space_after is not None:
        spacing.set(qn('w:after'), str(int(space_after * 20)))
    else:
        spacing.set(qn('w:after'), '0')

    spacing.set(qn('w:line'), str(int(240 * line_spacing)))
    spacing.set(qn('w:lineRule'), 'auto')

    pPr.append(spacing)


def add_title(doc, title: str):
    """Add centred document title."""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(HEADING_2_SIZE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, space_after=0)

    # Skip two lines
    for _ in range(2):
        blank = doc.add_paragraph()
        set_paragraph_spacing(blank, space_before=0, space_after=0)


def add_recipient_block(doc, name: str, title: str, org: str, address: str, city: str, country: str):
    """Add recipient address block."""
    # Name (bold)
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(BODY_FONT_SIZE)
    set_paragraph_spacing(p, space_before=0, space_after=0)

    # Other lines
    for line in [title, org, address, city, country]:
        if line and line.strip():
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = FONT_NAME
            run.font.size = Pt(BODY_FONT_SIZE)
            set_paragraph_spacing(p, space_before=0, space_after=0)


def add_date(doc, date_str: str):
    """Add date with spacing."""
    for _ in range(2):
        blank = doc.add_paragraph()
        set_paragraph_spacing(blank, space_before=0, space_after=0)

    p = doc.add_paragraph()
    run = p.add_run(date_str)
    run.font.name = FONT_NAME
    run.font.size = Pt(BODY_FONT_SIZE)
    set_paragraph_spacing(p, space_before=0, space_after=0)


def add_salutation(doc, name: str):
    """Add salutation with spacing."""
    for _ in range(2):
        blank = doc.add_paragraph()
        set_paragraph_spacing(blank, space_before=0, space_after=0)

    p = doc.add_paragraph()
    run = p.add_run(f"Dear {name}")
    run.font.name = FONT_NAME
    run.font.size = Pt(BODY_FONT_SIZE)
    set_paragraph_spacing(p, space_before=0, space_after=0)


def add_body_paragraph(doc, text: str, is_first=False):
    """Add body paragraph with correct spacing."""
    p = doc.add_paragraph()

    # Handle bold text marked with **
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.name = FONT_NAME
        run.font.size = Pt(BODY_FONT_SIZE)

    set_paragraph_spacing(p, space_before=0, space_after=6)


def add_signoff(doc):
    """Add sign-off block."""
    blank = doc.add_paragraph()
    set_paragraph_spacing(blank, space_before=0, space_after=0)

    p = doc.add_paragraph()
    run = p.add_run("Kind regards,")
    run.font.name = FONT_NAME
    run.font.size = Pt(BODY_FONT_SIZE)
    set_paragraph_spacing(p, space_before=0, space_after=0)

    for _ in range(2):
        blank = doc.add_paragraph()
        set_paragraph_spacing(blank, space_before=0, space_after=0)

    # Name
    p = doc.add_paragraph()
    run = p.add_run(YOUR_NAME)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(BODY_FONT_SIZE)
    set_paragraph_spacing(p, space_before=0, space_after=0)

    # Qualifications
    p = doc.add_paragraph()
    run = p.add_run(YOUR_QUALIFICATIONS)
    run.font.name = FONT_NAME
    run.font.size = Pt(QUALIFICATIONS_FONT_SIZE)
    set_paragraph_spacing(p, space_before=0, space_after=0)

    # Title
    p = doc.add_paragraph()
    run = p.add_run(YOUR_TITLE)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(BODY_FONT_SIZE)
    set_paragraph_spacing(p, space_before=0, space_after=0)


def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()


def add_heading(doc, text: str, level: int = 2):
    """Add heading with space before."""
    text = re.sub(r'^#+\s*', '', text)

    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME

    if level == 2:
        run.font.size = Pt(HEADING_2_SIZE)
    elif level == 3:
        run.font.size = Pt(HEADING_3_SIZE)
    else:
        run.font.size = Pt(BODY_FONT_SIZE)

    set_paragraph_spacing(p, space_before=12, space_after=6)


def add_bullet_item(doc, text: str):
    """Add bullet list item with no space before/after."""
    text = re.sub(r'^[-*]\s*', '', text.strip())

    p = doc.add_paragraph()

    run = p.add_run('• ')
    run.font.name = FONT_NAME
    run.font.size = Pt(BODY_FONT_SIZE)

    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.name = FONT_NAME
        run.font.size = Pt(BODY_FONT_SIZE)

    p.paragraph_format.left_indent = Inches(0.25)
    set_paragraph_spacing(p, space_before=0, space_after=0)


def add_numbered_item(doc, text: str, number: int):
    """Add numbered list item."""
    text = re.sub(r'^\d+\.\s*', '', text.strip())

    p = doc.add_paragraph()

    run = p.add_run(f'{number}. ')
    run.font.name = FONT_NAME
    run.font.size = Pt(BODY_FONT_SIZE)

    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.name = FONT_NAME
        run.font.size = Pt(BODY_FONT_SIZE)

    p.paragraph_format.left_indent = Inches(0.25)
    set_paragraph_spacing(p, space_before=0, space_after=0)


def add_nested_bullet_item(doc, text: str, indent_level: int = 1):
    """Add nested bullet list item."""
    text = re.sub(r'^[-*]\s*', '', text.strip())

    p = doc.add_paragraph()

    bullet = '○' if indent_level > 0 else '•'
    run = p.add_run(f'{bullet} ')
    run.font.name = FONT_NAME
    run.font.size = Pt(BODY_FONT_SIZE)

    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.name = FONT_NAME
        run.font.size = Pt(BODY_FONT_SIZE)

    p.paragraph_format.left_indent = Inches(0.25 + (0.25 * indent_level))
    set_paragraph_spacing(p, space_before=0, space_after=0)


def process_content_sections(doc, content: str):
    """Process markdown content sections into Word elements."""
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        raw_line = lines[i]
        line = raw_line.strip()

        if not line or line == '---':
            i += 1
            continue

        indent = len(raw_line) - len(raw_line.lstrip())

        if line.startswith('###'):
            add_heading(doc, line, level=3)
        elif line.startswith('##'):
            add_heading(doc, line, level=2)
        elif line.startswith('#'):
            add_heading(doc, line, level=1)
        elif indent >= 2 and (line.startswith('- ') or line.startswith('* ')):
            indent_level = indent // 2
            add_nested_bullet_item(doc, line, indent_level)
        elif line.startswith('- ') or line.startswith('* '):
            add_bullet_item(doc, line)
        elif re.match(r'^\d+\.', line):
            num = int(re.match(r'^(\d+)\.', line).group(1))
            add_numbered_item(doc, line, num)
        elif line.startswith('**') and ':' in line:
            add_body_paragraph(doc, line)
        else:
            add_body_paragraph(doc, line)

        i += 1


def extract_body_paragraphs(cover_letter: str) -> list:
    """Extract body paragraphs from cover letter markdown."""
    lines = cover_letter.split('\n')
    paragraphs = []
    in_body = False
    current_para = []

    for line in lines:
        if line.startswith('#') or line.startswith('**') or line.startswith('---'):
            if current_para:
                paragraphs.append(' '.join(current_para))
                current_para = []
            continue

        if line.strip().startswith('Dear '):
            in_body = True
            continue

        if 'Kind regards' in line or line.strip() == 'Kind regards,':
            if current_para:
                paragraphs.append(' '.join(current_para))
            break

        if in_body:
            if line.strip():
                current_para.append(line.strip())
            elif current_para:
                paragraphs.append(' '.join(current_para))
                current_para = []

    return paragraphs


def create_document(
    template_path: Path,
    source_path: Path,
    output_path: Path,
    recipient_name: str,
    recipient_title: str,
    recipient_org: str,
    recipient_address: str,
    recipient_city: str,
    recipient_country: str,
    doc_title: str,
    date_str: str
):
    """Create the Word document."""
    doc = Document(template_path)

    # Clear body content but preserve headers/footers (stored in sectPr)
    body = doc.element.body

    for element in list(body):
        if element.tag != qn('w:sectPr'):
            body.remove(element)

    sections = parse_markdown(source_path)

    # Add cover letter
    add_title(doc, doc_title)
    add_recipient_block(doc, recipient_name, recipient_title, recipient_org,
                        recipient_address, recipient_city, recipient_country)
    add_date(doc, date_str)
    add_salutation(doc, recipient_name)

    body_paras = extract_body_paragraphs(sections['cover_letter'])
    for i, para in enumerate(body_paras):
        add_body_paragraph(doc, para, is_first=(i == 0))

    add_signoff(doc)

    if sections['content_sections']:
        add_page_break(doc)
        process_content_sections(doc, sections['content_sections'])

    doc.save(output_path)
    print(f"Document saved to: {output_path}")


def main():
    parser = argparse.ArgumentParser(description='Create professional client documents from markdown')
    parser.add_argument('--template', required=True, help='Path to letterhead template (.docx)')
    parser.add_argument('--source', required=True, help='Path to markdown source file')
    parser.add_argument('--output', required=True, help='Output path for Word document')
    parser.add_argument('--recipient-name', required=True, help='Recipient name')
    parser.add_argument('--recipient-title', default='', help='Recipient title/position')
    parser.add_argument('--recipient-org', default='', help='Recipient organisation')
    parser.add_argument('--recipient-address', default='', help='Street address')
    parser.add_argument('--recipient-city', default='', help='City and postcode')
    parser.add_argument('--recipient-country', default='Ireland', help='Country')
    parser.add_argument('--doc-title', required=True, help='Document title')
    parser.add_argument('--date', default='', help='Date string (defaults to today)')

    args = parser.parse_args()

    if not args.date:
        from datetime import datetime
        args.date = datetime.now().strftime('%A %d %B %Y')

    create_document(
        template_path=Path(args.template),
        source_path=Path(args.source),
        output_path=Path(args.output),
        recipient_name=args.recipient_name,
        recipient_title=args.recipient_title,
        recipient_org=args.recipient_org,
        recipient_address=args.recipient_address,
        recipient_city=args.recipient_city,
        recipient_country=args.recipient_country,
        doc_title=args.doc_title,
        date_str=args.date
    )


if __name__ == '__main__':
    main()
